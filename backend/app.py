# ---------------------------
# backend/app.py
# ---------------------------
import os
import platform
import subprocess
import urllib.parse
import tempfile
import shutil
from pathlib import Path
from typing import Optional, List, Dict, Iterator
from datetime import datetime

from fastapi import FastAPI, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from .config import (
    ALLOWED_ROOTS,
    DEFAULT_PAGE_SIZE,
    MAX_PAGE_SIZE,
    USE_INDEX,
    DB_PATH,
    CORS_ALLOW_ORIGINS,
    PARENT_ORDER,
    # NEW: report config
    MULTI_EXP_LIMIT,
    REPORTS_OUTPUT_DIR,
    DOCX_REPORT_BASENAME,
    DOCX_REPORT_AUTHOR,
)
from .search import walk_search, coverage_rows
from .search import monthly_coverage  # monthly data
from .search import multi_exp_missing  # NEW: multi-EXP summary
from .indexer import search_index  # optional
from .mime_types import guess_type

APP_DIR = Path(__file__).resolve().parent
FRONTEND_DIR = APP_DIR.parent / "frontend"

app = FastAPI(title="Smart Document Finder", version="2.2.0")

# ---- CORS ----
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ALLOW_ORIGINS or ["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---- Health ----
@app.get("/health")
def health():
    return {"status": "ok"}


# ---- Helpers ----
def _path_allowed(p: Path) -> bool:
    """
    Ensure the requested path is within one of the allowed roots.
    Works across different path casing and UNC/drive forms.
    """
    try:
        rp = os.path.normcase(os.path.abspath(str(p)))
    except Exception:
        return False

    for root in ALLOWED_ROOTS:
        try:
            base = os.path.normcase(os.path.abspath(root))
            if os.path.commonpath([rp, base]) == base:
                return True
        except Exception:
            continue
    return False


def _dedup_key(path_str: str) -> str:
    """
    Build a robust identity key for a path: prefer (st_dev, st_ino),
    fall back to normalized absolute path.
    """
    p = Path(path_str)
    try:
        st = p.stat()
        return f"{st.st_dev}:{st.st_ino}"
    except Exception:
        try:
            return os.path.normcase(os.path.abspath(str(p)))
        except Exception:
            return str(p).lower()


def _fmt_size(n: Optional[int]) -> Optional[str]:
    if n is None:
        return None
    try:
        n = int(n)
    except Exception:
        return None
    units = ["B", "KB", "MB", "GB", "TB"]
    i = 0
    f = float(n)
    while f >= 1024 and i < len(units) - 1:
        f /= 1024.0
        i += 1
    return f"{f:.1f} {units[i]}"


def _fmt_mtime(ts: Optional[float]) -> Optional[str]:
    try:
        return datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M")
    except Exception:
        return None


def _has_children(dir_path: Path) -> bool:
    try:
        with os.scandir(dir_path) as it:
            for entry in it:
                try:
                    if entry.is_dir(follow_symlinks=False):
                        return True
                except PermissionError:
                    continue
    except PermissionError:
        return False
    return False


# =========================================================
#  A) BASELINE SEARCH
# =========================================================
@app.get("/search")
def search(
    query: str = Query(..., min_length=1, description="Filename-only, case-insensitive substring"),
    year: Optional[str] = Query(None),
    month: Optional[str] = Query(None),
    company: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(DEFAULT_PAGE_SIZE, ge=1, le=MAX_PAGE_SIZE),
):
    if USE_INDEX:
        res = search_index(
            db_path=str(DB_PATH),
            query=query,
            year=year,
            month=month,
            company=company,
            page=page,
            page_size=page_size,
        )
        total = res.get("count", 0)
        total_pages = res.get("total_pages", 1)
        items = res.get("items", [])
    else:
        res = walk_search(
            query=query,
            roots=ALLOWED_ROOTS,
            year=year,
            month=month,
            company=company,
            page=page,
            page_size=page_size,
        )
        total = res.get("count", 0)
        total_pages = res.get("total_pages", 1)
        items = res.get("items", [])

    return JSONResponse({
        "count": total,
        "page": page,
        "page_size": page_size,
        "total_pages": total_pages,
        "items": items,
    })


# =========================================================
#  B) COVERAGE ROWS (single table: exactly 7 parents)
# =========================================================
@app.get("/coverage-rows")
def coverage_rows_endpoint(
    query: str = Query(..., min_length=1),
    year: Optional[str] = Query(None),
    month: Optional[str] = Query(None),
    company: Optional[str] = Query(None),
):
    res = coverage_rows(
        parents_order=PARENT_ORDER,
        roots=ALLOWED_ROOTS,
        query=query,
        year=year,
        month=month,
        company=company,
    )
    return JSONResponse(res)


@app.get("/coverage-files")
def coverage_files_endpoint(
    parent: str = Query(..., description="One of the fixed parent names (e.g., 'CIPL')"),
    query: str = Query(..., min_length=1),
    year: Optional[str] = Query(None),
    month: Optional[str] = Query(None),
    company: Optional[str] = Query(None),
):
    """
    For a given parent (e.g., 'PDIR REPORT'), return *files* for the child table:
      - If a search hit is a FILE, include it directly.
      - If a search hit is a FOLDER (e.g., 'EXP-192'), enumerate files inside it
        so Preview works. (Expands up to max_depth and returns files.)
    """
    res = walk_search(
        query=query or "",
        roots=ALLOWED_ROOTS,
        year=year,
        month=month,
        company=company,
        page=1,
        page_size=10_000,
    )

    wanted = (parent or "").strip().lower()
    items_out: List[Dict] = []

    for item in res.get("items", []):
        kind = (item.get("kind") or "file").lower()
        full_path = Path(item.get("full_path", ""))

        parent_name = full_path.parent.name.lower()

        # Only keep entries belonging to the requested parent
        if parent_name != wanted:
            if wanted not in [p.lower() for p in full_path.parts]:
                continue

        if kind == "file":
            items_out.append({
                "kind": "file",
                "file_name": item.get("file_name") or full_path.name,
                "parent_folder": full_path.parent.name,
                "full_path": str(full_path),
            })
        else:
            if not _path_allowed(full_path):
                continue
            items_out.extend(_list_files_in_folder(full_path, max_depth=2, limit=500))

    # De-duplicate by robust file identity (device+inode) or normalized path
    seen = set()
    deduped = []
    for it in items_out:
        key = _dedup_key(it["full_path"])
        if key in seen:
            continue
        seen.add(key)
        deduped.append(it)

    return {"items": deduped}


# =========================================================
#  B2) MONTHLY COVERAGE (Available vs Missing by parent)
# =========================================================
@app.get("/monthly-coverage")
def monthly_coverage_endpoint(
    year: str = Query(..., description="Year (e.g., 2025)"),
    month: str = Query(..., description="Full month name (e.g., August)"),
    company: Optional[str] = Query(None),
):
    """
    Returns grouped results for a selected Year+Month:
      {
        "period": {"year":"2025","month":"august"},
        "available": [ { "parent":"CIPL","count":N,"items":[...] }, ... ],
        "missing":   [ "DOCK AUDIT REPORT", "BL", ... ]
      }
    """
    res = monthly_coverage(
        parents_order=PARENT_ORDER,
        roots=ALLOWED_ROOTS,
        year=year,
        month=month,
        company=company,
        max_items_per_parent=5000,
    )
    return JSONResponse(res)


# =========================================================
#  C) PREVIEW / BASIC BROWSE
# =========================================================
@app.get("/preview")
def preview(path: str):
    p = Path(path)
    if not _path_allowed(p):
        raise HTTPException(status_code=403, detail="Path not allowed")
    if not p.exists():
        raise HTTPException(status_code=404, detail="File not found")

    media_type = guess_type(str(p))
    headers = {"Content-Disposition": f'inline; filename="{p.name}"'}
    return FileResponse(str(p), media_type=media_type, headers=headers)


@app.get("/browse")
def browse(path: str, query: Optional[str] = None):
    """
    JSON API for programmatic listing (immediate children).
    """
    p = Path(path)
    if not _path_allowed(p):
        raise HTTPException(status_code=403, detail="Path not allowed")
    if not p.exists() or not p.is_dir():
        raise HTTPException(status_code=404, detail="Folder not found")

    ql = (query or "").lower()
    items = []
    try:
        for child in p.iterdir():
            name = child.name
            if ql and ql not in name.lower():
                continue
            items.append({
                "kind": "folder" if child.is_dir() else "file",
                "file_name": name,
                "parent_folder": p.name,
                "full_path": str(child),
            })
    except PermissionError:
        raise HTTPException(status_code=403, detail="Permission denied when listing folder")

    return {"count": len(items), "items": items}


# =========================================================
#  D) RICH BROWSER ENDPOINTS (tree + details)
# =========================================================
@app.get("/browse-children")
def browse_children(path: str = Query(...)):
    """
    Return immediate *folders* of a path, with has_children flag for each.
    """
    base = Path(path)
    if not _path_allowed(base):
        raise HTTPException(status_code=403, detail="Path not allowed")
    if not base.exists() or not base.is_dir():
        raise HTTPException(status_code=404, detail="Folder not found")

    out = []
    try:
        with os.scandir(base) as it:
            for entry in it:
                try:
                    if entry.is_dir(follow_symlinks=False):
                        fp = Path(entry.path)
                        out.append({
                            "name": entry.name,
                            "full_path": str(fp),
                            "has_children": _has_children(fp),
                        })
                except PermissionError:
                    continue
    except PermissionError:
        raise HTTPException(status_code=403, detail="Permission denied when listing folder")

    out.sort(key=lambda r: r["name"].lower())
    return {"items": out}


def _iter_items(root: Path, include_subfolders: bool) -> Iterator[Path]:
    if not include_subfolders:
        try:
            with os.scandir(root) as it:
                for entry in it:
                    yield Path(entry.path)
        except PermissionError:
            return
        return

    # deep
    for dirpath, dirnames, filenames in os.walk(root):
        pdir = Path(dirpath)
        for d in dirnames:
            yield pdir / d
        for f in filenames:
            yield pdir / f


@app.get("/browse-list")
def browse_list(
    path: str = Query(...),
    q: Optional[str] = Query(None),
    include_subfolders: bool = Query(False),
    offset: int = Query(0, ge=0),
    limit: int = Query(1000, ge=1, le=5000),
):
    """
    Return folders+files for the details pane (optionally recursive), with soft pagination.
    """
    base = Path(path)
    if not _path_allowed(base):
        raise HTTPException(status_code=403, detail="Path not allowed")
    if not base.exists() or not base.is_dir():
        raise HTTPException(status_code=404, detail="Folder not found")

    ql = (q or "").lower()
    rows: List[Dict] = []

    max_collect = offset + limit + 1
    count_collected = 0

    def push(p: Path):
        nonlocal count_collected
        name = p.name
        if ql and ql not in name.lower():
            return
        count_collected += 1
        if count_collected <= offset:
            return
        if len(rows) >= limit + 1:
            return

        try:
            st = p.stat()
        except PermissionError:
            st = None

        rows.append({
            "name": name,
            "full_path": str(p),
            "kind": "folder" if p.is_dir() else "file",
            "size": _fmt_size(st.st_size) if st and p.is_file() else None,
            "modified": _fmt_mtime(st.st_mtime) if st else None,
        })

    collected = 0
    for child in _iter_items(base, include_subfolders):
        push(child)
        collected += 1
        if collected >= max_collect:
            break

    rows.sort(key=lambda r: (r["kind"] != "folder", (r["name"] or "").lower()))

    has_more = len(rows) > limit
    if has_more:
        rows = rows[:limit]
    next_offset = (offset + limit) if has_more else None

    return {"items": rows, "has_more": has_more, "next_offset": next_offset}


# ---- Server-rendered folder UI (two-pane, with sidebar toggle) ----
@app.get("/browse-ui", response_class=HTMLResponse)
def browse_ui(path: str = Query(..., description="Absolute path under one of ALLOWED_ROOTS")):
    """
    Two-pane browser (tree + details). Uses /browse-children and /browse-list.
    Raw HTML string with placeholders replaced.
    """
    p = Path(path)
    if not _path_allowed(p):
        raise HTTPException(status_code=403, detail="Path not allowed")
    if not p.exists():
        raise HTTPException(status_code=404, detail="Not found")
    if p.is_file():
        media_type = guess_type(str(p))
        headers = {"Content-Disposition": f'inline; filename="{p.name}"'}
        return FileResponse(str(p), media_type=media_type, headers=headers)

    allowed_roots_encoded = urllib.parse.quote("|".join(ALLOWED_ROOTS))
    initial_path_encoded = urllib.parse.quote(str(p))

    # (HTML omitted here for brevity in this comment — it is identical to your latest working version)
    html = r"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Browse</title>
  <meta name="viewport" content="width=device-width,initial-scale=1" />
  <style>
    :root {
      --bg:#0b1020; --bg-2:#0e1630; --card:#111a33; --border:#1e2a44;
      --text:#e8eef7; --muted:#9fb3cc; --accent:#22d3ee; --accent-strong:#06b6d4;
      --success:#10b981; --danger:#ef4444;
    }
    *{box-sizing:border-box} html,body{height:100%}
    body{ margin:0; background:linear-gradient(180deg,var(--bg),var(--bg-2)); color:var(--text);
      font-family:ui-sans-serif,-apple-system,Segoe UI,Roboto,Helvetica,Arial; }
    .wrap{display:grid; grid-template-columns:0 1fr; height:100vh; transition:grid-template-columns .2s ease}
    .wrap.sidebar-open{grid-template-columns:320px 1fr}
    .sidebar{border-right:1px solid var(--border); background:rgba(15,23,42,.5); padding:12px; overflow:auto}
    .main{display:flex; flex-direction:column; min-width:0}
    .topbar{ display:flex; align-items:center; justify-content:space-between;
      padding:12px 16px; border-bottom:1px solid var(--border);
      background:linear-gradient(180deg,rgba(20,28,52,.75),rgba(16,24,43,.65)); position:sticky; top:0; z-index:20; }
    .crumbs a{color:#cfe2ff; text-decoration:none} .crumbs a:hover{text-decoration:underline}
    .controls{display:flex; gap:10px; align-items:center}
    .input,.checkbox{background:linear-gradient(180deg,#0c1736,#0c1530); border:1px solid #203258; color:var(--text); border-radius:10px; padding:8px 10px}
    .btn{padding:8px 10px; border-radius:10px; border:1px solid #203258; background:#0f1f3d; color:var(--text); cursor:pointer}
    .btn:hover{background:#162a52}
    .iconbtn{width:36px; height:36px; display:inline-grid; place-items:center; border-radius:10px; border:1px solid #203258; background:rgba(255,255,255,.06); cursor:pointer}
    .iconbtn:hover{background:rgba(255,255,255,.10)}
    .tree ul{list-style:none; margin:0; padding-left:16px} .tree li{margin:3px 0}
    .node{display:flex; align-items:center; gap:6px; cursor:pointer}
    .twisty{width:12px; height:12px; display:inline-block; border-left:6px solid #8fb3ff; border-top:6px solid transparent; border-bottom:6px solid transparent}
    .node.expanded .twisty{transform:rotate(90deg)}
    .content{padding:10px 16px; overflow:auto}
    table{width:100%; border-collapse:separate; border-spacing:0; border:1px solid var(--border); border-radius:12px; overflow:hidden}
    thead th{text-align:left; padding:10px; background:linear-gradient(180deg,#0f1c3e,#0c1733); color:#d7e7ff; border-bottom:1px solid var(--border)}
    tbody td{padding:10px; border-bottom:1px solid var(--border)}
    tbody tr:nth-child(odd) td{background:#0b1530} tbody tr:nth-child(even) td{background:#0d1a3c}
    tbody tr:hover td{background:#12224a}
    .muted{color:#9fb3cc} .empty{color:#cbd5e1; padding:14px; border:1px dashed var(--border); border-radius:12px; background:linear-gradient(180deg,#0e1a38,#0c1633)}
  </style>
</head>
<body>
  <div class="wrap">
    <aside class="sidebar"><div class="tree" id="tree"></div></aside>
    <main class="main">
      <div class="topbar">
        <div class="crumbs" id="crumbs">Loading…</div>
        <div class="controls">
          <button id="btnSidebar" class="iconbtn" title="Open sidebar" aria-label="Open sidebar">
            <svg viewBox="0 0 24 24" fill="#cfe2ff">
              <rect x="3" y="4" width="18" height="16" rx="2" ry="2" fill="none" stroke="#8fb3ff" stroke-width="2"></rect>
              <rect x="3" y="4" width="6" height="16" rx="2" ry="2" fill="#8fb3ff" opacity=".35"></rect>
            </svg>
          </button>
          <input class="input" id="search" placeholder="Search in this folder…" />
          <label class="grid-row">
            <input type="checkbox" id="deep" class="checkbox" />
            <span class="muted">Include subfolders</span>
          </label>
        </div>
      </div>
      <div class="content">
        <table>
          <thead><tr><th>Name</th><th class="nowrap">Type</th><th class="nowrap">Size</th><th class="nowrap">Modified</th><th class="nowrap">Open</th></tr></thead>
          <tbody id="rows"><tr><td colspan="5" class="muted">Loading…</td></tr></tbody>
        </table>
        <div id="moreWrap" style="padding:10px 0;display:none;">
          <button class="btn" id="btnMore">Load more</button>
          <span class="muted" id="moreInfo"></span>
        </div>
      </div>
    </main>
  </div>
  <script>
  /* (JS identical to your last good version — omitted in this snippet for brevity) */
  </script>
</body>
</html>
"""
    html = html.replace("__ALLOWED_ROOTS__", allowed_roots_encoded).replace("__INITIAL_PATH__", initial_path_encoded)
    return HTMLResponse(html)


# ---- Optional zip (kept; not shown in UI now) ----
@app.get("/zip-folder")
def zip_folder(path: str = Query(..., description="Absolute folder path under allowed roots")):
    folder = Path(path)
    if not _path_allowed(folder):
        raise HTTPException(status_code=403, detail="Path not allowed")
    if not folder.exists() or not folder.is_dir():
        raise HTTPException(status_code=404, detail="Folder not found")

    tmpdir = Path(tempfile.mkdtemp())
    zip_base = tmpdir / folder.name
    shutil.make_archive(str(zip_base), "zip", folder)
    return FileResponse(str(zip_base) + ".zip", filename=f"{folder.name}.zip")


# ---- Windows shell-open endpoints (compat) ----
class ShellRequest(BaseModel):
    path: str


@app.post("/shell/open")
def shell_open(req: ShellRequest):
    if platform.system().lower() != "windows":
        raise HTTPException(status_code=501, detail="Open supported only on Windows")
    p = Path(req.path)
    if not _path_allowed(p):
        raise HTTPException(status_code=403, detail="Path not allowed")
    if not p.exists():
        raise HTTPException(status_code=404, detail="File not found")
    try:
        os.startfile(str(p))
        return {"status": "launched", "target": str(p)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to open: {e}")


@app.post("/shell/open-folder")
def shell_open_folder(req: ShellRequest):
    if platform.system().lower() != "windows":
        raise HTTPException(status_code=501, detail="Open supported only on Windows")
    p = Path(req.path)
    if not _path_allowed(p):
        raise HTTPException(status_code=403, detail="Path not allowed")

    target = p if p.is_dir() else p.parent
    if not target.exists():
        raise HTTPException(status_code=404, detail="Folder not found")

    try:
        if p.is_file():
            subprocess.run(["explorer", f"/select,{str(p)}"], check=False)
        else:
            subprocess.run(["explorer", str(target)], check=False)
        return {"status": "launched", "target": str(target)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to open folder: {e}")


# === Multi-EXP Missing: API endpoints ===
class MultiMissingRequest(BaseModel):
    exps: List[str]
    year: Optional[str] = None
    month: Optional[str] = None
    company: Optional[str] = None

@app.get("/multi-missing")
def multi_missing_get(
    exps: str = Query(..., description="Comma-separated EXPs like EXP-123,EXP-124"),
    year: Optional[str] = None,
    month: Optional[str] = None,
    company: Optional[str] = None,
):
    exp_list = [e.strip() for e in (exps or "").split(",") if e.strip()]
    out = multi_exp_missing(
        parents_order=PARENT_ORDER,
        roots=ALLOWED_ROOTS,
        exp_codes=exp_list,
        year=year,
        month=month,
        company=company,
        limit=MULTI_EXP_LIMIT,
    )
    return JSONResponse({
        "results": out.get("items", []),
        "summary": out.get("summary", {}),
        "invalid": out.get("invalid", []),
        "limit": out.get("limit", MULTI_EXP_LIMIT),
    })

@app.post("/multi-missing-docx")
def multi_missing_docx(req: MultiMissingRequest):
    out = multi_exp_missing(
        parents_order=PARENT_ORDER,
        roots=ALLOWED_ROOTS,
        exp_codes=req.exps or [],
        year=req.year,
        month=req.month,
        company=req.company,
        limit=MULTI_EXP_LIMIT,
    )

    # Try to create a .docx; fall back to .txt if python-docx is unavailable
    try:
        from docx import Document
        doc = Document()
        doc.core_properties.author = DOCX_REPORT_AUTHOR
        doc.add_heading('Missing Folders Report', level=1)

        meta = doc.add_paragraph()
        meta.add_run(f"Year: {req.year or '—'}   Month: {req.month or '—'}   Company: {req.company or '—'}")

        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "EXP"
        hdr[1].text = "Missing Folders"

        for item in out.get("items", []):
            row = table.add_row().cells
            row[0].text = item.get("exp", "")
            row[1].text = ", ".join(item.get("missing", [])) or "—"

        REPORTS_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        tmp = REPORTS_OUTPUT_DIR / f"{DOCX_REPORT_BASENAME}.docx"
        doc.save(tmp)
        return FileResponse(str(tmp), filename=f"{DOCX_REPORT_BASENAME}.docx")

    except Exception:
        tmpdir = Path(tempfile.mkdtemp())
        txt = tmpdir / (DOCX_REPORT_BASENAME + ".txt")
        with open(txt, "w", encoding="utf-8") as f:
          f.write("Missing Folders Report\n")
          f.write(f"Year: {req.year or '—'}  Month: {req.month or '—'}  Company: {req.company or '—'}\n\n")
          for item in out.get("items", []):
              miss = ", ".join(item.get("missing", [])) or "—"
              f.write(f"{item.get('exp','')}: {miss}\n")
        return FileResponse(str(txt), filename=f"{DOCX_REPORT_BASENAME}.txt")


# ---- Static frontend at ROOT (/) ----
if FRONTEND_DIR.exists():
    print("Serving frontend from:", FRONTEND_DIR.resolve())
    app.mount("/", StaticFiles(directory=FRONTEND_DIR, html=True), name="frontend")


# ---- Helper: list files inside a folder (used by /coverage-files) ----
def _list_files_in_folder(folder_path: Path, max_depth: int = 2, limit: int = 500):
    out = []
    if not folder_path.exists() or not folder_path.is_dir():
        return out

    base_parts = Path(folder_path).resolve().parts
    for dirpath, _dirnames, filenames in os.walk(folder_path):
        depth = len(Path(dirpath).resolve().parts) - len(base_parts)
        if depth > max_depth:
            continue
        for fname in filenames:
            full = Path(dirpath) / fname
            out.append({
                "kind": "file",
                "file_name": fname,
                "parent_folder": Path(dirpath).name,
                "full_path": str(full),
            })
            if len(out) >= limit:
                return out
    return out
